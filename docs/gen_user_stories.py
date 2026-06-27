from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x29, 0x5C)   # dark navy
TEAL   = RGBColor(0x00, 0x7A, 0x87)   # teal accent
LGRAY  = RGBColor(0xF4, 0xF6, 0xF8)   # light gray background
DGRAY  = RGBColor(0x4A, 0x4A, 0x4A)   # dark gray body text
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GOLD   = RGBColor(0xC8, 0x96, 0x20)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'FFFFFF'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run_with_style(para, text, bold=False, italic=False,
                       size=11, color=DGRAY, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return run

def heading(text, level=1, color=NAVY, size=None, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True):
    sizes = {1: 20, 2: 15, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size if size else sizes[level])
    run.font.color.rgb = color
    return p

def body(text, indent=0, space_after=6, italic=False, color=DGRAY, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p

def bullet(text, indent=0.2, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = DGRAY
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '007A87')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(cover, '0A295C')
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after  = Pt(4)
r = cover.add_run('WAH4H')
r.font.name  = 'Calibri'
r.font.size  = Pt(42)
r.font.bold  = True
r.font.color.rgb = WHITE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub, '0A295C')
sub.paragraph_format.space_after = Pt(4)
r2 = sub.add_run('We Are Health for Hospitals')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(18)
r2.font.color.rgb = RGBColor(0xC8, 0x96, 0x20)
r2.bold = True

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(tagline, '0A295C')
tagline.paragraph_format.space_after = Pt(40)
r3 = tagline.add_run('Hospital Information System')
r3.font.name  = 'Calibri'
r3.font.size  = Pt(13)
r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

spacer = doc.add_paragraph()
shade_paragraph(spacer, '0A295C')
spacer.paragraph_format.space_after = Pt(2)

doc_title = doc.add_paragraph()
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(doc_title, '0A295C')
doc_title.paragraph_format.space_after = Pt(4)
rt = doc_title.add_run('USER STORIES AND SYSTEM GUIDEBOOK')
rt.font.name  = 'Calibri'
rt.font.size  = Pt(16)
rt.font.bold  = True
rt.font.color.rgb = WHITE

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(ver, '0A295C')
ver.paragraph_format.space_after = Pt(60)
rv = ver.add_run('Version 1.0   |   June 2026')
rv.font.name  = 'Calibri'
rv.font.size  = Pt(11)
rv.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1  INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading('1.  About This Document', level=1)
divider()
body(
    'This guidebook describes the WAH4H Hospital Information System from the perspective '
    'of the people who use it every day. Each section of this document presents a system '
    'module, explains what it does, and lists the user stories that define how staff members '
    'interact with it. The acceptance criteria for each story serve as the agreed definition '
    'of "done" between the development team and hospital stakeholders.'
)
body(
    'WAH4H is built for Philippine Local Government Unit (LGU) hospitals. It covers the '
    'full patient journey from admission to discharge, including clinical documentation, '
    'laboratory services, pharmacy management, billing, and PhilHealth claims processing. '
    'The system is also connected to the WAH4PC inter-hospital data exchange network, which '
    'allows patient records to be shared securely between participating facilities.'
)

heading('2.  System Overview', level=1)
divider()
body('The following diagram shows the major modules and how they connect to each other along the patient care workflow.')

# Module flow table
tbl = doc.add_table(rows=1, cols=9)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
flow_labels = [
    'Registration', 'Admission', 'Monitoring', 'Laboratory',
    'Pharmacy', 'Billing', 'PhilHealth', 'Discharge', 'Gateway'
]
flow_colors = [
    '007A87', '0A295C', '2E7D32', '6A1B9A',
    'BF360C', 'C62828', 'AD1457', '4E342E', '00838F'
]
for i, cell in enumerate(tbl.rows[0].cells):
    shade_cell(cell, flow_colors[i])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(flow_labels[i])
    run.font.name  = 'Calibri'
    run.font.size  = Pt(8)
    run.font.bold  = True
    run.font.color.rgb = WHITE

doc.add_paragraph().paragraph_format.space_after = Pt(8)

heading('3.  User Roles', level=1)
divider()
body('WAH4H supports six distinct staff roles. Access to each module depends on the role assigned to a user account.')

roles = [
    ('System Administrator',
     'Manages the system configuration, user accounts, hospital locations, practitioner '
     'profiles, and inter-hospital gateway settings. Has full access to all modules.'),
    ('Doctor',
     'Registers patients, creates and closes encounters, documents procedures and diagnoses, '
     'orders laboratory tests, writes prescriptions, and finalizes discharge summaries.'),
    ('Nurse',
     'Admits patients, assigns ward and bed locations, records vital signs and clinical '
     'observations, and administers medications as directed by the attending physician.'),
    ('Pharmacist',
     'Manages the medication inventory, processes prescription requests from doctors, '
     'dispenses medications, and monitors stock levels and expiry dates.'),
    ('Laboratory Technician',
     'Accepts laboratory orders from doctors, manages specimen collection, enters test '
     'results, handles imaging study records, and maintains the laboratory test catalog.'),
    ('Billing Officer',
     'Creates billing accounts for patient encounters, generates invoices, submits PhilHealth '
     'insurance claims, records payments, and manages payment reconciliation.'),
]

tbl2 = doc.add_table(rows=len(roles) + 1, cols=2)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl2.columns[0].width = Inches(1.8)
tbl2.columns[1].width = Inches(4.2)

hdr_cells = tbl2.rows[0].cells
for cell, txt in zip(hdr_cells, ['Role', 'Responsibilities']):
    shade_cell(cell, '0A295C')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.05)
    run = p.add_run(txt)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = WHITE

row_colors = ['F4F6F8', 'FFFFFF']
for i, (role, desc) in enumerate(roles):
    row = tbl2.rows[i + 1]
    shade_cell(row.cells[0], row_colors[i % 2])
    shade_cell(row.cells[1], row_colors[i % 2])
    for cell, txt, bold in zip(row.cells, [role, desc], [True, False]):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Inches(0.05)
        run = p.add_run(txt)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(9.5)
        run.font.bold  = bold
        run.font.color.rgb = DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(4)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# MODULE TEMPLATE FUNCTION
# ════════════════════════════════════════════════════════════════════════════
def module_banner(title, icon_text, color_hex, description, who_uses, key_features):
    p = doc.add_paragraph()
    shade_paragraph(p, color_hex)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f'  {icon_text}  {title}')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = WHITE

    sub = doc.add_paragraph()
    shade_paragraph(sub, color_hex)
    sub.paragraph_format.space_after = Pt(12)
    sub.paragraph_format.left_indent = Inches(0.1)
    rs = sub.add_run(description)
    rs.font.name  = 'Calibri'
    rs.font.size  = Pt(10)
    rs.font.color.rgb = RGBColor(0xDD, 0xEE, 0xFF)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    c_who  = tbl.rows[0].cells[0]
    c_feat = tbl.rows[0].cells[1]
    shade_cell(c_who,  'EEF3F8')
    shade_cell(c_feat, 'EEF3F8')

    pw = c_who.paragraphs[0]
    shade_cell(c_who, 'E8F0FE')
    pw.paragraph_format.space_before = Pt(6)
    rw = pw.add_run('Who Uses This Module')
    rw.font.name = 'Calibri'
    rw.font.size = Pt(9)
    rw.font.bold = True
    rw.font.color.rgb = NAVY

    for u in who_uses:
        p2 = c_who.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run(f'  {u}')
        r2.font.name  = 'Calibri'
        r2.font.size  = Pt(9)
        r2.font.color.rgb = DGRAY

    pf = c_feat.paragraphs[0]
    shade_cell(c_feat, 'E8F0FE')
    pf.paragraph_format.space_before = Pt(6)
    rf = pf.add_run('Key Features at a Glance')
    rf.font.name = 'Calibri'
    rf.font.size = Pt(9)
    rf.font.bold = True
    rf.font.color.rgb = NAVY

    for f in key_features:
        p3 = c_feat.add_paragraph()
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after  = Pt(2)
        r3 = p3.add_run(f'  {f}')
        r3.font.name  = 'Calibri'
        r3.font.size  = Pt(9)
        r3.font.color.rgb = DGRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def story_block(story_id, title, role, intent, value, criteria, priority, points):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell = tbl.rows[0].cells[0]
    shade_cell(cell, 'F9FAFB')

    hdr = cell.paragraphs[0]
    hdr.paragraph_format.space_before = Pt(6)
    hdr.paragraph_format.left_indent  = Inches(0.1)
    id_run = hdr.add_run(story_id + '  ')
    id_run.font.name  = 'Calibri'
    id_run.font.size  = Pt(9)
    id_run.font.bold  = True
    id_run.font.color.rgb = TEAL
    title_run = hdr.add_run(title)
    title_run.font.name  = 'Calibri'
    title_run.font.size  = Pt(11)
    title_run.font.bold  = True
    title_run.font.color.rgb = NAVY

    p_meta = cell.add_paragraph()
    p_meta.paragraph_format.left_indent = Inches(0.1)
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta = p_meta.add_run(f'Priority: {priority}   |   Story Points: {points}')
    r_meta.font.name  = 'Calibri'
    r_meta.font.size  = Pt(8.5)
    r_meta.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r_meta.italic = True

    story_text = (
        f'As a {role}, I want to {intent} so that {value}.'
    )
    ps = cell.add_paragraph()
    ps.paragraph_format.left_indent = Inches(0.1)
    ps.paragraph_format.space_after = Pt(6)
    rs = ps.add_run(story_text)
    rs.font.name  = 'Calibri'
    rs.font.size  = Pt(10)
    rs.font.color.rgb = DGRAY
    rs.italic = True

    pc = cell.add_paragraph()
    pc.paragraph_format.left_indent = Inches(0.1)
    pc.paragraph_format.space_after = Pt(3)
    rc = pc.add_run('Acceptance Criteria')
    rc.font.name  = 'Calibri'
    rc.font.size  = Pt(9.5)
    rc.font.bold  = True
    rc.font.color.rgb = NAVY

    for c in criteria:
        pb = cell.add_paragraph()
        pb.paragraph_format.left_indent = Inches(0.2)
        pb.paragraph_format.space_after = Pt(2)
        rb = pb.add_run(f'☐  {c}')
        rb.font.name  = 'Calibri'
        rb.font.size  = Pt(9)
        rb.font.color.rgb = DGRAY

    cell.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4  MODULE 1 AUTH
# ════════════════════════════════════════════════════════════════════════════
heading('4.  Authentication and Account Management', level=1)
divider()
module_banner(
    title='Authentication and Account Management',
    icon_text='[AUTH]',
    color_hex='0A295C',
    description=(
        'Controls who can access the system. Every staff member must have a verified account '
        'before using any module. Logins are protected by two-step verification.'
    ),
    who_uses=['All hospital staff', 'System Administrator'],
    key_features=[
        'Two-step registration with email OTP verification',
        'Secure login with JWT access and refresh tokens',
        'Password reset via email OTP',
        'Role-based access control across all modules',
        'Session timeout after 15 minutes of inactivity',
    ]
)

story_block('AUTH-001', 'Initiate Staff Account Registration',
    role='new hospital staff member',
    intent='submit my registration details and receive a one-time password on my email',
    value='I can begin the process of creating a verified account in WAH4H',
    criteria=[
        'The registration form collects first name, last name, email, password, role, and organization.',
        'On form submission, the system validates that the email is not already registered.',
        'A six-digit OTP is sent to the provided email address.',
        'The form shows a clear error if the email is already in use.',
        'The password must be at least 12 characters, not purely numeric, and not a common password.',
    ],
    priority='High', points=5
)

story_block('AUTH-002', 'Complete Registration via OTP Verification',
    role='new hospital staff member who received an OTP',
    intent='enter my OTP to verify my email and activate my account',
    value='I can log in and start using the system',
    criteria=[
        'The OTP input screen is displayed after initiation.',
        'The system accepts the correct OTP and creates the user account.',
        'An incorrect OTP shows a clear error message.',
        'The OTP expires after a defined time window and the user can request a new one.',
        'On success, the user is redirected to the login screen.',
    ],
    priority='High', points=3
)

story_block('AUTH-003', 'Initiate Two-Step Login',
    role='registered staff member',
    intent='log in using my email and password',
    value='the system can verify my identity before granting access',
    criteria=[
        'The login form accepts email and password.',
        'On correct credentials, an OTP is sent to the registered email.',
        'On incorrect credentials, a descriptive error is shown without revealing account existence.',
        'Login attempts are rate limited to a maximum of 60 per minute per IP address.',
    ],
    priority='High', points=3
)

story_block('AUTH-004', 'Complete Login via OTP Verification',
    role='staff member who received a login OTP',
    intent='enter my OTP to complete login',
    value='I receive my access and refresh tokens and can use the system',
    criteria=[
        'A correct OTP returns a 15-minute access token and a 7-day refresh token.',
        'An incorrect OTP shows an error without exposing token data.',
        'Successful login redirects the user to the dashboard for their assigned role.',
        'Tokens are stored in secure browser storage.',
    ],
    priority='High', points=3
)

story_block('AUTH-005', 'Refresh JWT Access Token',
    role='logged-in staff member',
    intent='have my session automatically refreshed',
    value='I am not abruptly logged out during active work',
    criteria=[
        'The frontend silently calls the token refresh endpoint before the access token expires.',
        'A valid refresh token returns a new access token.',
        'An expired or blacklisted refresh token redirects the user to the login screen.',
        'The old refresh token is blacklisted after rotation.',
    ],
    priority='High', points=3
)

story_block('AUTH-006', 'Initiate Password Reset',
    role='staff member who has forgotten their password',
    intent='request a password reset via my registered email',
    value='I can regain access to my account',
    criteria=[
        'A Forgot Password link is accessible from the login page.',
        'Submitting a registered email sends a password reset OTP to that address.',
        'Password reset requests are rate limited to a maximum of 3 per minute.',
        'The response does not reveal whether the email exists in the system.',
    ],
    priority='High', points=3
)

story_block('AUTH-007', 'Complete Password Reset via OTP',
    role='staff member who received a password reset OTP',
    intent='enter the OTP and set a new password',
    value='I can log in again with my new credentials',
    criteria=[
        'The OTP and new password are submitted together.',
        'The new password must meet OWASP-compliant rules: minimum 12 characters, not numeric only, and not a common password.',
        'On success, the old password is invalidated and the user is redirected to login.',
        'Expired or invalid OTP shows a descriptive error.',
    ],
    priority='High', points=3
)

story_block('AUTH-008', 'Change Password While Logged In',
    role='logged-in staff member',
    intent='change my password from my account settings',
    value='I can maintain account security proactively',
    criteria=[
        'The change password flow requires an OTP verification step.',
        'Current password confirmation is required before the OTP is sent.',
        'The new password must meet OWASP-compliant rules.',
        'On success, all existing tokens are invalidated and the user is redirected to login.',
    ],
    priority='Medium', points=3
)

story_block('AUTH-009', 'Role-Based Access Control',
    role='system administrator',
    intent='assign a specific role to each staff account',
    value='each user only sees and interacts with the modules relevant to their job',
    criteria=[
        'Available roles are: System Administrator, Doctor, Nurse, Lab Technician, Pharmacist, and Billing Officer.',
        'Each role grants access to a defined set of modules.',
        'Navigation menus display only the modules the current user role can access.',
        'API endpoints return HTTP 403 Forbidden when a user attempts to access a restricted resource.',
        'Role assignment is visible and editable by administrators only.',
    ],
    priority='High', points=8
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# MODULE 2  PATIENTS
# ════════════════════════════════════════════════════════════════════════════
heading('5.  Patient Management', level=1)
divider()
module_banner(
    title='Patient Management',
    icon_text='[PAT]',
    color_hex='007A87',
    description=(
        'The central record for every patient in the facility. All other modules reference '
        'the patient record. Supports PhilHealth ID validation, allergy tracking, '
        'immunization history, and inter-hospital data exchange via the WAH4PC gateway.'
    ),
    who_uses=['Doctor', 'Nurse'],
    key_features=[
        'Patient registration with demographic and contact information',
        'PhilHealth ID format validation',
        'Allergy and drug intolerance tracking',
        'Immunization history recording',
        'Search by name, PhilHealth ID, or system ID',
        'Fetch or push patient records to the WAH4PC inter-hospital network',
    ]
)

story_block('PAT-001', 'Register a New Patient',
    role='doctor or nurse',
    intent='register a new patient with their demographic and contact information',
    value='the patient has a record in the system that other modules can reference',
    criteria=[
        'Registration form collects first name, last name, date of birth, sex, civil status, nationality, address, and contact number.',
        'Optional fields include PhilHealth ID (validated format XX-XXXXXXXXX-X), blood type, PWD status, and indigenous people status.',
        'Emergency contact information can be recorded.',
        'Duplicate PhilHealth IDs are rejected with a clear error.',
        'The saved patient appears in the patient list with a unique system ID.',
    ],
    priority='High', points=8
)

story_block('PAT-002', 'Search and View Patient Records',
    role='doctor, nurse, or authorized staff member',
    intent='search for a patient by name, PhilHealth ID, or system ID',
    value='I can quickly retrieve their record to view or update it',
    criteria=[
        'A search bar is available on the patient list page.',
        'Search is not case sensitive and supports partial name matches.',
        'Results list shows name, date of birth, PhilHealth ID, and last encounter date.',
        'Clicking a patient opens their full profile.',
        'Results are paginated at 50 records per page.',
    ],
    priority='High', points=5
)

story_block('PAT-003', 'Edit Patient Demographics',
    role='doctor or nurse',
    intent='update a patient\'s demographic information',
    value='the record remains accurate when details change such as address or contact number',
    criteria=[
        'All demographic fields are editable from the patient profile.',
        'PhilHealth ID format is re-validated on edit.',
        'Changes are saved with an updated timestamp.',
        'An audit trail is maintained through the updated_at field.',
    ],
    priority='Medium', points=3
)

story_block('PAT-004', 'Record a Patient Condition or Diagnosis',
    role='doctor',
    intent='record a medical condition or diagnosis for a patient',
    value='it becomes part of their permanent clinical history',
    criteria=[
        'Condition form collects condition name or code, clinical status (active, resolved, or recurrence), severity, and onset date.',
        'Optional fields include abatement date, stage, evidence, and notes.',
        'A patient can have multiple conditions listed.',
        'Conditions are displayed on the patient profile in chronological order.',
        'Clinical status can be updated as the patient\'s condition changes.',
    ],
    priority='High', points=5
)

story_block('PAT-005', 'Record Patient Allergies and Intolerances',
    role='doctor or nurse',
    intent='record a patient\'s known allergies or drug intolerances',
    value='clinicians and pharmacists are warned before administering substances',
    criteria=[
        'Allergy form collects substance, clinical status, and criticality (low, high, or unable to assess).',
        'Reaction details can be added including manifestation and severity.',
        'Multiple allergies per patient are supported.',
        'The allergy list is prominently visible on the patient profile.',
        'Pharmacists see allergy warnings when processing prescriptions.',
    ],
    priority='High', points=5
)

story_block('PAT-006', 'Record Patient Immunization History',
    role='doctor or nurse',
    intent='record a patient\'s vaccinations',
    value='their immunization history is available for clinical decisions',
    criteria=[
        'Immunization form collects vaccine code or name, date administered, lot number, and expiry date.',
        'The administering practitioner can be linked.',
        'Protocol and series information is recordable such as dose 1 of 3.',
        'Multiple immunization records per patient are supported.',
        'Records are sorted by date administered.',
    ],
    priority='Medium', points=5
)

story_block('PAT-007', 'Fetch Patient Records from the WAH4PC Gateway',
    role='doctor or nurse',
    intent='query the WAH4PC network for a patient\'s records from other hospitals',
    value='I have a complete clinical picture even if the patient was previously treated elsewhere',
    criteria=[
        'A Fetch from Gateway action is available on the patient profile.',
        'The system sends a query to the WAH4PC gateway using the patient\'s PhilHealth ID.',
        'Returned data including encounters, procedures, and immunizations is displayed for review before import.',
        'The transaction is logged in the WAH4PC transaction audit table with type pull and status tracked.',
        'Network errors are handled gracefully with a user-friendly message.',
    ],
    priority='High', points=8
)

story_block('PAT-008', 'Push Patient Records to the WAH4PC Gateway',
    role='doctor',
    intent='send a patient\'s records to the WAH4PC network',
    value='other hospitals treating the same patient can access their history',
    criteria=[
        'A Send to Gateway action is available on the patient profile.',
        'The system packages the patient\'s record as a FHIR bundle and sends it to the gateway.',
        'The transaction is logged in the WAH4PC transaction audit table with type push and status tracked.',
        'Success and failure states are communicated to the user.',
        'Only users with the Doctor role can trigger a push.',
    ],
    priority='High', points=8
)

page_break()

# MODULE 3 ADM
heading('6.  Admission and Encounters', level=1)
divider()
module_banner(
    title='Admission and Encounters',
    icon_text='[ADM]',
    color_hex='1A5276',
    description=(
        'An encounter is the primary event that links every clinical activity together. '
        'When a patient arrives at the hospital, an encounter is created. All procedures, '
        'observations, lab orders, and medications are recorded against that encounter.'
    ),
    who_uses=['Doctor', 'Nurse', 'System Administrator'],
    key_features=[
        'Create inpatient, outpatient, and emergency encounter records',
        'Assign patients to specific wards, rooms, and beds',
        'Record hospitalization details including diet and special arrangements',
        'Document clinical procedures during an encounter',
        'View and filter all active encounters',
        'Close encounters to trigger the discharge workflow',
    ]
)

story_block('ADM-001', 'Create a New Encounter',
    role='doctor or nurse',
    intent='create an encounter record when a patient arrives at the hospital',
    value='all clinical activities are linked to this visit',
    criteria=[
        'Encounter form requires patient, encounter class (inpatient, outpatient, or emergency), and start date and time.',
        'Optional fields include admission source and expected discharge date.',
        'The encounter is linked to the selected patient record.',
        'A newly created encounter appears in the list with the status in-progress.',
        'Doctor, nurse, and administrator roles can create encounters.',
    ],
    priority='High', points=5
)

story_block('ADM-002', 'Assign Patient to a Ward, Room, and Bed',
    role='nurse',
    intent='assign an admitted patient to a specific ward, room, and bed',
    value='I know the patient\'s physical location for care delivery',
    criteria=[
        'The location picker shows the hierarchy: Building, Wing, Ward, Room, and Bed.',
        'Only beds with available status are selectable.',
        'Assigning a bed marks it as occupied.',
        'The assigned location is displayed on the encounter record.',
        'Changing the location updates bed status accordingly.',
    ],
    priority='High', points=5
)

story_block('ADM-003', 'Record Hospitalization Details',
    role='nurse or doctor',
    intent='record the details of a patient\'s hospitalization including diet and special arrangements',
    value='the care team can accommodate the patient\'s needs throughout their stay',
    criteria=[
        'Diet preference field accepts free text or coded values.',
        'Special courtesy and special arrangement fields are available.',
        'A re-admission checkbox flags returning patients.',
        'These details are saved on the encounter and visible to all authorized staff.',
    ],
    priority='Medium', points=3
)

story_block('ADM-004', 'Document a Procedure During an Encounter',
    role='doctor',
    intent='record a medical procedure performed on a patient during their encounter',
    value='the procedure is part of the clinical record and can be referenced in billing',
    criteria=[
        'Procedure form requires procedure code or name, status, patient, and encounter reference.',
        'Optional fields include performer, body site, outcome, complications, and follow-up recommendations.',
        'Multiple procedures can be recorded per encounter.',
        'Procedures are listed on the encounter detail page.',
        'A completed procedure can be referenced when building a billing claim.',
    ],
    priority='High', points=5
)

story_block('ADM-005', 'View and Manage Active Encounters',
    role='nurse or doctor',
    intent='see a list of all currently active encounters',
    value='I can monitor which patients are admitted and access their records quickly',
    criteria=[
        'Encounter list shows patient name, encounter class, location, admission date, and attending physician.',
        'Filters by encounter class, location, date range, and status are available.',
        'Clicking an encounter opens the full encounter detail page.',
        'Encounters are sorted by admission date with the most recent shown first.',
        'Search by patient name is supported.',
    ],
    priority='High', points=5
)

story_block('ADM-006', 'Close an Encounter',
    role='doctor',
    intent='mark an encounter as finished when the patient has been discharged',
    value='it is no longer listed as active and the patient\'s visit is properly closed',
    criteria=[
        'An End Encounter action is available on the encounter detail page.',
        'The end date and time is recorded automatically.',
        'The encounter status changes to finished.',
        'Closed encounters remain accessible for historical reference.',
        'Closing an encounter triggers the discharge workflow if no discharge record has been created.',
    ],
    priority='High', points=3
)

page_break()

# MODULE 4 MON
heading('7.  Monitoring and Vital Signs', level=1)
divider()
module_banner(
    title='Monitoring and Vital Signs',
    icon_text='[MON]',
    color_hex='2E7D32',
    description=(
        'Tracks patient health indicators throughout their admission. Nurses record '
        'measurements such as temperature, heart rate, blood pressure, and oxygen saturation. '
        'Abnormal values are flagged automatically to support early clinical response.'
    ),
    who_uses=['Nurse', 'Doctor'],
    key_features=[
        'Record individual vital sign observations with FHIR coded types',
        'Multi-component blood pressure entry (systolic and diastolic)',
        'Chronological observation timeline per encounter',
        'Automatic flagging of values outside reference ranges',
        'Charge item creation for billable monitoring services',
    ]
)

story_block('MON-001', 'Record a Patient Vital Sign Observation',
    role='nurse',
    intent='record a vital sign measurement for a patient',
    value='the patient\'s condition is tracked over time during their admission',
    criteria=[
        'Observation form requires type, value, unit, effective date and time, patient, and encounter.',
        'Common vital sign types are available via a coded picker using FHIR observation codes.',
        'The recorded value is saved against the patient\'s current encounter.',
        'The observation appears in the patient\'s monitoring timeline.',
    ],
    priority='High', points=5
)

story_block('MON-002', 'Record a Multi-Component Observation for Blood Pressure',
    role='nurse',
    intent='record blood pressure as a single observation with separate systolic and diastolic values',
    value='both components are stored and displayed together for clinical accuracy',
    criteria=[
        'Blood pressure entry has two sub-fields: systolic in mmHg and diastolic in mmHg.',
        'Both components are saved as child records linked to the parent observation.',
        'Display shows the format 120/80 mmHg in the monitoring timeline.',
        'Reference ranges are applied independently to each component.',
        'Abnormal component values are flagged individually.',
    ],
    priority='High', points=5
)

story_block('MON-003', 'View Patient Observation History',
    role='doctor or nurse',
    intent='view the historical trend of a patient\'s vital signs during their current encounter',
    value='I can assess their progress and identify deterioration early',
    criteria=[
        'The monitoring page shows a timeline of all observations for the selected encounter.',
        'Each entry shows observation type, value, unit, date and time, and recorded by.',
        'Observations can be filtered by type and date range.',
        'The timeline is sorted with the most recent observation shown first.',
    ],
    priority='High', points=5
)

story_block('MON-004', 'Flag Abnormal Observation Values',
    role='nurse or doctor',
    intent='see automatic flagging when a recorded value falls outside the normal reference range',
    value='I can respond to potentially dangerous readings quickly',
    criteria=[
        'Each observation type has a configurable reference range with minimum and maximum values.',
        'Values outside the range are highlighted visually using color coding: yellow for borderline and red for critical.',
        'An interpretation code such as H for high, L for low, or N for normal is stored on the observation.',
        'Abnormal observations can be filtered to the top of the list.',
    ],
    priority='High', points=5
)

story_block('MON-005', 'Create a Charge Item for a Billable Monitoring Service',
    role='nurse',
    intent='create a charge item when performing a billable monitoring service such as ECG monitoring',
    value='the service is captured in the patient\'s billing account',
    criteria=[
        'Charge item form requires service code, quantity, unit price, patient, and encounter.',
        'The charge item is linked to the active billing account for the encounter.',
        'Created charge items appear in the patient\'s billing summary.',
        'Performer and enterer fields can be recorded for accountability.',
    ],
    priority='Medium', points=3
)

page_break()

# MODULE 5 LAB
heading('8.  Laboratory', level=1)
divider()
module_banner(
    title='Laboratory',
    icon_text='[LAB]',
    color_hex='6A1B9A',
    description=(
        'Manages the full laboratory workflow from test ordering by doctors through specimen '
        'collection, result entry, and final result delivery. Also tracks imaging studies '
        'such as X-ray and CT scan records.'
    ),
    who_uses=['Doctor', 'Nurse', 'Laboratory Technician'],
    key_features=[
        'Lab test ordering with priority levels: routine, urgent, and STAT',
        'Specimen collection and chain of custody documentation',
        'Result entry with reference range flagging',
        'Laboratory test catalog management with pricing',
        'Imaging study tracking for radiology records',
        'Results visible to ordering doctors and nurses',
    ]
)

story_block('LAB-001', 'Order a Lab Test for a Patient',
    role='doctor',
    intent='place a lab test order for a patient',
    value='the lab technician can collect specimens and perform the test',
    criteria=[
        'Lab order form requires test type from catalog, patient, encounter, and requester.',
        'Priority can be set to routine, urgent, or STAT.',
        'The created diagnostic report has status registered.',
        'The lab order appears in the lab queue visible to lab technicians.',
        'A charge item is automatically created for the ordered test.',
    ],
    priority='High', points=5
)

story_block('LAB-002', 'Record Specimen Collection',
    role='lab technician',
    intent='record the collection of a specimen for a lab order',
    value='the sample\'s chain of custody is documented',
    criteria=[
        'Specimen form collects collection method, body site, collection date and time, and received time.',
        'The specimen is linked to the corresponding diagnostic report.',
        'A notes field is available for special collection circumstances.',
        'The specimen record updates the diagnostic report status to in-progress.',
    ],
    priority='High', points=3
)

story_block('LAB-003', 'Enter and Finalize Lab Test Results',
    role='lab technician',
    intent='enter the results of a completed lab test',
    value='doctors can review the findings and make clinical decisions',
    criteria=[
        'Result entry is available on each diagnostic report with status in-progress.',
        'Results support numeric values, text, and coded values.',
        'Finalizing the report changes its status to final and records the issued date and time.',
        'Completed results are visible to the ordering doctor and nurse.',
    ],
    priority='High', points=8
)

story_block('LAB-004', 'Manage the Lab Test Catalog',
    role='lab technician or system administrator',
    intent='manage the catalog of available lab tests',
    value='doctors can order from an up-to-date list with correct pricing and reference ranges',
    criteria=[
        'Catalog entry includes test name, code, category, base price, and turnaround time.',
        'Reference ranges are configurable per test with support for age-based variations.',
        'New tests can be added and existing tests can be edited or deactivated.',
        'Deactivated tests no longer appear in the order form\'s test picker.',
        'Price changes are reflected in new charge items while existing items remain unaffected.',
    ],
    priority='Medium', points=5
)

story_block('LAB-005', 'Track Imaging Studies',
    role='lab technician or doctor',
    intent='record imaging studies such as X-ray, CT scan, and ultrasound for a patient',
    value='the imaging history is part of the patient\'s clinical record',
    criteria=[
        'Imaging study form collects modality, start date and time, and number of series and instances.',
        'The interpreter can be linked to the study.',
        'The study is associated with the patient\'s encounter.',
        'Imaging studies are listed separately from laboratory test reports.',
        'Records support attaching notes or findings text.',
    ],
    priority='Medium', points=5
)

story_block('LAB-006', 'View Lab Results as a Doctor or Nurse',
    role='doctor or nurse',
    intent='view the finalized lab results for my patient',
    value='I can incorporate the findings into my clinical assessment',
    criteria=[
        'Lab results are accessible from the patient\'s encounter page.',
        'Finalized reports show test name, all result values, reference ranges, and flags for out-of-range values.',
        'Results are sorted by collection date with the most recent shown first.',
        'Pending and in-progress reports show their current status.',
        'STAT and urgent orders are visually distinguished from routine orders.',
    ],
    priority='High', points=3
)

page_break()

# MODULE 6 PHARMACY
heading('9.  Pharmacy', level=1)
divider()
module_banner(
    title='Pharmacy',
    icon_text='[PHR]',
    color_hex='BF360C',
    description=(
        'Covers the complete medication lifecycle: prescription by doctors, dispensing by '
        'pharmacists, administration by nurses, and inventory management. Patient allergy '
        'alerts are shown at the point of prescribing.'
    ),
    who_uses=['Doctor', 'Pharmacist', 'Nurse'],
    key_features=[
        'Medication prescription with dosage and dispense instructions',
        'Allergy warnings at point of prescribing',
        'Prescription queue with priority sorting',
        'Medication dispensing and administration records',
        'Inventory management with stock adjustments',
        'Low-stock alerts and expiry date tracking by batch number',
    ]
)

story_block('PHR-001', 'Prescribe Medication for a Patient',
    role='doctor',
    intent='create a medication prescription for a patient',
    value='the pharmacist can dispense the correct medication',
    criteria=[
        'Prescription form requires medication from catalog, patient, encounter, intent, and priority.',
        'Dosage instructions include dose quantity, route, frequency, and timing.',
        'Dispense instructions include validity period, quantity to dispense, and refills allowed.',
        'The created prescription appears in the pharmacy queue with status active.',
        'Patient known allergies are shown as warnings during prescription.',
    ],
    priority='High', points=8
)

story_block('PHR-002', 'Dispense Medication to a Patient',
    role='pharmacist or nurse',
    intent='record that a medication has been dispensed and administered to a patient',
    value='the administration is documented and inventory is updated',
    criteria=[
        'A medication administration record is created referencing the original prescription.',
        'Administration records the medication given, dose, route, date and time, and administering practitioner.',
        'The prescription status is updated to completed after the final dose.',
        'Inventory stock count is decremented by the dispensed quantity.',
        'A charge item is created for the medication against the patient\'s billing account.',
    ],
    priority='High', points=8
)

story_block('PHR-003', 'View and Manage Medication Inventory',
    role='pharmacist',
    intent='view the current medication inventory and manage stock levels',
    value='I can ensure medications are available when needed and prevent stockouts',
    criteria=[
        'Inventory list shows medication name, current stock, unit, reorder level, expiry date, batch number, and unit cost.',
        'Stock can be adjusted manually for receipts, returns, and corrections.',
        'Expired medications are flagged visually.',
        'Inventory is searchable by medication name or code.',
        'Stock movements are recorded with a reason and timestamp.',
    ],
    priority='High', points=5
)

story_block('PHR-004', 'Receive Low-Stock Alerts',
    role='pharmacist',
    intent='be notified when a medication\'s stock falls below its reorder level',
    value='I can reorder in time to avoid a stockout',
    criteria=[
        'Each inventory item has a configurable reorder level.',
        'When stock falls at or below the reorder level, the item is flagged in the inventory list.',
        'A summary of low-stock medications is visible on the pharmacy dashboard.',
        'The alert is cleared when stock is replenished above the reorder level.',
    ],
    priority='Medium', points=3
)

story_block('PHR-005', 'Track Medication Batch Numbers and Expiry Dates',
    role='pharmacist',
    intent='record the batch number and expiry date for each medication stock entry',
    value='I can identify and quarantine affected stock in case of a product recall',
    criteria=[
        'Batch number and expiry date fields are present on inventory entries.',
        'Medications approaching expiry are highlighted based on a configurable threshold.',
        'Expired medications are clearly marked and excluded from dispensing selections.',
        'Batch number is recorded on each medication administration for traceability.',
    ],
    priority='Medium', points=3
)

story_block('PHR-006', 'View the Prescriptions Queue',
    role='pharmacist',
    intent='view all active prescriptions awaiting dispensing',
    value='I can work through them in priority order',
    criteria=[
        'The prescription queue shows patient name, medication, dose, route, priority, ordering doctor, and time ordered.',
        'Queue is filterable by priority (STAT, urgent, routine) and date.',
        'STAT prescriptions appear at the top by default.',
        'Each prescription can be opened for full detail and action.',
        'Already dispensed prescriptions are moved to a completed view.',
    ],
    priority='High', points=5
)

page_break()

# MODULE 7 BILLING
heading('10.  Billing and Financial Management', level=1)
divider()
module_banner(
    title='Billing and Financial Management',
    icon_text='[BIL]',
    color_hex='C62828',
    description=(
        'Accumulates charge items from all clinical modules into a unified billing account '
        'per encounter. Generates invoices, processes PhilHealth insurance claims, records '
        'payments, and manages payment reconciliation.'
    ),
    who_uses=['Billing Officer', 'Doctor', 'System Administrator'],
    key_features=[
        'Billing account per encounter with automatic charge item accumulation',
        'Invoice generation with itemized cost breakdown',
        'PhilHealth insurance claim submission with ICD codes and procedures',
        'Payment recording with method and reference tracking',
        'Payment reconciliation against submitted claims',
        'Payment notice management',
    ]
)

story_block('BIL-001', 'Create a Billing Account for an Encounter',
    role='billing clerk',
    intent='create a billing account linked to a patient\'s encounter',
    value='all charge items generated during the stay are accumulated in one place',
    criteria=[
        'Billing account form requires patient, encounter, and service period start and end dates.',
        'Optional fields include guarantor and coverage reference for PhilHealth.',
        'A patient may have one active billing account per encounter.',
        'Charge items from labs, pharmacy, and monitoring link to this account automatically.',
    ],
    priority='High', points=5
)

story_block('BIL-002', 'Generate an Invoice for a Patient',
    role='billing clerk',
    intent='generate an invoice from a patient\'s accumulated charge items',
    value='the total cost of care is presented to the patient or guarantor',
    criteria=[
        'Invoice is generated from all charge items linked to the billing account.',
        'Invoice displays an itemized list of services, quantities, unit prices, subtotal, and total.',
        'Invoice can be marked as draft, issued, balanced, or cancelled.',
        'An issued invoice is accessible in printable format.',
        'Invoice date and due date are recorded.',
    ],
    priority='High', points=8
)

story_block('BIL-003', 'Submit a PhilHealth Insurance Claim',
    role='billing clerk or doctor',
    intent='submit a PhilHealth claim for a patient\'s encounter',
    value='the hospital can recover the covered cost from PhilHealth',
    criteria=[
        'Claim form collects claim type, patient, encounter, provider, insurer, billable period, and total amount.',
        'Diagnoses with ICD codes and procedures can be attached as claim line items.',
        'Care team members are listed on the claim.',
        'Submitted claim status is tracked as active, cancelled, or entered in error.',
        'Claim items include revenue code, service code, service date, quantity, and unit price.',
    ],
    priority='High', points=13
)

story_block('BIL-004', 'Record a Payment Against an Invoice',
    role='billing clerk',
    intent='record a payment made by the patient or their guarantor',
    value='the account balance is updated accurately',
    criteria=[
        'Payment form records amount, payment date, payment method, and reference number.',
        'Payment is linked to the patient\'s billing account and invoice.',
        'The outstanding balance on the account is recalculated after payment is recorded.',
        'Overpayments are flagged for review.',
        'Payment history is viewable on the billing account.',
    ],
    priority='High', points=5
)

story_block('BIL-005', 'Process Payment Reconciliation',
    role='billing clerk',
    intent='reconcile payments received against submitted claims',
    value='I can confirm that the correct amount was reimbursed by PhilHealth or other payers',
    criteria=[
        'Reconciliation record links a payment to one or more claim line items.',
        'Discrepancies between claimed and paid amounts are highlighted.',
        'Reconciled claim items are marked as settled.',
        'A reconciliation summary report is available per billing period.',
    ],
    priority='Medium', points=8
)

story_block('BIL-006', 'Send and View Payment Notices',
    role='billing clerk',
    intent='record payment notices sent to patients or payers',
    value='there is a documented record of payment communications',
    criteria=[
        'Payment notice form records recipient, payment reference, notice date, and status.',
        'Notices are linked to the relevant invoice or claim.',
        'A list of all payment notices is accessible and filterable by date and status.',
        'Notice status can be updated to reflect sent, acknowledged, or disputed.',
    ],
    priority='Low', points=3
)

page_break()

# MODULE 8 DISCHARGE
heading('11.  Discharge', level=1)
divider()
module_banner(
    title='Discharge',
    icon_text='[DIS]',
    color_hex='4E342E',
    description=(
        'Formalizes the end of a patient\'s hospital stay. The doctor completes a discharge '
        'summary, records follow-up instructions, and finalizes the record. Finalization '
        'closes the encounter and frees the patient\'s assigned bed.'
    ),
    who_uses=['Doctor', 'Nurse', 'Billing Officer'],
    key_features=[
        'Discharge summary creation linked to the patient encounter',
        'Free-text discharge instructions and follow-up plan',
        'Pending items tracking for outstanding results or referrals',
        'Finalization that closes the encounter and releases the bed',
        'Read-only discharge history for clinical and billing reference',
    ]
)

story_block('DIS-001', 'Create a Discharge Summary',
    role='doctor',
    intent='create a discharge summary for a patient who is leaving the hospital',
    value='the patient and any follow-up providers have a complete record of their stay',
    criteria=[
        'Discharge form requires patient, encounter, and discharge date and time.',
        'Summary of stay in free text must be completed.',
        'Form is linked to the patient\'s current encounter.',
        'The created discharge record shows status draft until finalized.',
    ],
    priority='High', points=5
)

story_block('DIS-002', 'Record Discharge Instructions and Follow-Up Plan',
    role='doctor',
    intent='add discharge instructions and a follow-up plan to the discharge summary',
    value='the patient knows what to do after leaving the hospital',
    criteria=[
        'Discharge instructions field accepts free text including medications, activity restrictions, wound care, and other details.',
        'Follow-up plan field records follow-up type, responsible provider, and timeframe.',
        'Pending items such as outstanding lab results and referrals can be listed.',
        'Instructions are accessible from the discharge record and printable for the patient.',
    ],
    priority='High', points=3
)

story_block('DIS-003', 'Finalize and Submit Discharge Record',
    role='doctor',
    intent='finalize the discharge record',
    value='the encounter is officially closed and the patient is marked as discharged',
    criteria=[
        'A Finalize Discharge action changes the discharge status from draft to final.',
        'Finalizing the discharge updates the associated encounter status to finished.',
        'The bed and location assigned to the patient is freed and marked as available.',
        'The discharge record becomes read-only after finalization.',
        'Only the creating doctor or an administrator can finalize the discharge.',
    ],
    priority='High', points=5
)

story_block('DIS-004', 'View Discharge History',
    role='doctor, nurse, or billing clerk',
    intent='view a list of completed discharges',
    value='I can reference past discharge summaries for clinical or billing purposes',
    criteria=[
        'Discharge list shows patient name, encounter, discharge date, discharging doctor, and status.',
        'List is filterable by date range, doctor, and status.',
        'Clicking a discharge record opens the full summary.',
        'Finalized discharge records cannot be edited.',
    ],
    priority='Medium', points=3
)

page_break()

# MODULE 9 SYS
heading('12.  System Administration', level=1)
divider()
module_banner(
    title='System Administration',
    icon_text='[SYS]',
    color_hex='37474F',
    description=(
        'Provides tools for administrators to configure the hospital environment before '
        'clinical staff begin using the system. Covers organizations, locations, '
        'practitioner profiles, roles, and inter-hospital gateway endpoints.'
    ),
    who_uses=['System Administrator'],
    key_features=[
        'Hospital organization and branch management',
        'Physical location hierarchy: Building, Wing, Ward, Room, and Bed',
        'Practitioner profile creation and management',
        'Role and availability schedule assignment',
        'WAH4PC gateway API endpoint configuration',
        'System health monitoring dashboard',
    ]
)

story_block('SYS-001', 'Manage Hospital Organizations',
    role='system administrator',
    intent='add and manage hospital and healthcare organization records',
    value='staff and patients can be associated with the correct institution',
    criteria=[
        'Organization form collects name, type, address, contact information, and license number.',
        'Organizations appear in staff registration and encounter forms.',
        'Existing organizations can be edited or deactivated.',
        'Deactivated organizations are hidden from selection dropdowns.',
    ],
    priority='High', points=3
)

story_block('SYS-002', 'Manage Hospital Locations',
    role='system administrator',
    intent='define and manage the physical locations within the hospital',
    value='patients can be accurately assigned to specific wards, rooms, and beds',
    criteria=[
        'Location hierarchy supports Building, Wing, Ward, Room, and Bed.',
        'Each location has a name, type, status (available, occupied, or under maintenance), and parent location.',
        'New locations can be added at any level of the hierarchy.',
        'Bed status updates automatically when patients are admitted or discharged.',
    ],
    priority='High', points=8
)

story_block('SYS-003', 'Manage Practitioner Profiles',
    role='system administrator',
    intent='create and manage practitioner staff profiles',
    value='doctors, nurses, and other staff can be linked to clinical activities',
    criteria=[
        'Practitioner form collects first name, last name, gender, date of birth, and professional license number.',
        'Practitioners can be assigned to an organization and department.',
        'A practitioner profile is created automatically on account registration.',
        'Profiles can be edited or deactivated by administrators.',
        'Deactivated practitioners are excluded from clinical assignment pickers.',
    ],
    priority='High', points=5
)

story_block('SYS-004', 'Manage Practitioner Roles and Availability',
    role='system administrator',
    intent='assign roles and availability schedules to practitioners',
    value='the system reflects who is on duty and what services they provide',
    criteria=[
        'Practitioner role record includes practitioner, role code, organization, and availability schedule.',
        'A practitioner can have multiple roles.',
        'Availability is expressed as recurring weekly schedules.',
        'Role assignments are visible when assigning practitioners to encounters and procedures.',
    ],
    priority='Medium', points=5
)

story_block('SYS-005', 'Configure API Endpoints for Interoperability',
    role='system administrator',
    intent='manage the API endpoint records used for inter-hospital communication',
    value='the WAH4PC gateway and other external integrations remain correctly configured',
    criteria=[
        'Endpoint record collects name, URL, connection type, and status (active or suspended).',
        'Endpoints are associated with the hospital organization.',
        'An endpoint can be tested from the admin interface.',
        'Only administrators can create or modify endpoint records.',
    ],
    priority='Medium', points=3
)

story_block('SYS-006', 'Monitor System Health',
    role='system administrator',
    intent='view a system health report',
    value='I can proactively identify and resolve infrastructure issues',
    criteria=[
        'Health check endpoint returns database connectivity, cache status, and external service reachability.',
        'A liveness probe endpoint returns a simple OK response.',
        'Health status is accessible from the admin dashboard.',
        'Unhealthy services are highlighted with a descriptive error.',
    ],
    priority='High', points=3
)

page_break()

# MODULE 10 DASH + SEC
heading('13.  Dashboard and Analytics', level=1)
divider()
module_banner(
    title='Dashboard and Analytics',
    icon_text='[DASH]',
    color_hex='00695C',
    description=(
        'Provides each role with a tailored home screen showing the metrics and pending '
        'tasks most relevant to their work. Administrators see a hospital-wide view. '
        'Clinical staff see their own workload at a glance.'
    ),
    who_uses=['All roles'],
    key_features=[
        'Admin dashboard: total patients, active encounters, today\'s admissions and discharges',
        'Role-specific dashboards tailored to doctor, nurse, pharmacist, lab technician, and billing officer workloads',
        'Real-time bed occupancy view by ward',
        'Pending task counters per department',
    ]
)

story_block('DASH-001', 'View the System-Wide Administrative Dashboard',
    role='system administrator',
    intent='see a consolidated overview of the hospital\'s key metrics',
    value='I can make informed operational decisions',
    criteria=[
        'Dashboard displays total patients registered, active encounters, today\'s admissions, and today\'s discharges.',
        'Bed occupancy rate is shown as a percentage and broken down by ward.',
        'Pending items per department including lab orders, prescriptions, and billing invoices are summarized.',
        'Dashboard auto-refreshes without requiring a full page reload.',
    ],
    priority='High', points=8
)

story_block('DASH-002', 'View a Role-Specific Dashboard',
    role='doctor, nurse, pharmacist, lab technician, or billing clerk',
    intent='see a dashboard tailored to my role',
    value='I can immediately see my pending tasks and relevant metrics without navigating the full system',
    criteria=[
        'Doctor dashboard shows patients admitted today, pending lab results to review, and prescriptions ordered.',
        'Nurse dashboard shows patients in assigned ward, pending medication administrations, and vital signs due.',
        'Pharmacist dashboard shows prescriptions awaiting dispensing, low-stock medications, and expired medications.',
        'Lab Technician dashboard shows pending lab orders, specimens collected today, and overdue results.',
        'Billing Clerk dashboard shows unbilled encounters, pending claims, and outstanding invoices.',
    ],
    priority='High', points=13
)

story_block('DASH-003', 'View Patient Census and Bed Occupancy',
    role='nurse or administrator',
    intent='see a real-time view of bed occupancy by ward',
    value='I can manage patient placement and communicate capacity to admissions',
    criteria=[
        'Ward-level view shows each bed\'s status: occupied, available, or under maintenance.',
        'Occupied beds show the patient\'s name and admission date.',
        'Total and available bed count is shown per ward.',
        'The view refreshes when bed assignments change on admission or discharge.',
    ],
    priority='Medium', points=8
)

page_break()

heading('14.  Security and Compliance', level=1)
divider()
module_banner(
    title='Security and Compliance',
    icon_text='[SEC]',
    color_hex='4527A0',
    description=(
        'Enforces security policies that protect patient data and system integrity. '
        'Covers password rules, rate limiting, session management, and an immutable '
        'audit log for all inter-hospital data exchanges.'
    ),
    who_uses=['System Administrator', 'All roles'],
    key_features=[
        'OWASP-compliant password policy enforced at registration, reset, and change',
        'Rate limiting on sensitive endpoints to prevent automated attacks',
        'Automatic session timeout after 15 minutes of inactivity',
        'Immutable audit log for all WAH4PC gateway transactions',
        'Account settings and profile management with OTP verification',
    ]
)

story_block('SEC-001', 'Update Account Profile and Personal Settings',
    role='logged-in staff member',
    intent='update my profile information and personal preferences',
    value='my account details remain accurate',
    criteria=[
        'Account settings page allows editing of display name and contact email.',
        'Profile changes are saved immediately with a success notification.',
        'Email changes trigger a re-verification OTP flow before taking effect.',
    ],
    priority='Low', points=3
)

story_block('SEC-002', 'Enforce OWASP-Compliant Password Policy',
    role='system administrator',
    intent='enforce a strong password policy for all user accounts',
    value='the system is protected against brute-force and credential stuffing attacks',
    criteria=[
        'Minimum password length is 12 characters.',
        'Passwords that are purely numeric are rejected.',
        'Common and dictionary passwords are rejected.',
        'These rules are enforced at registration, password reset, and password change.',
        'Clear error messages explain which rule was violated.',
    ],
    priority='High', points=3
)

story_block('SEC-003', 'Rate Limit Sensitive API Endpoints',
    role='system administrator',
    intent='have rate limiting applied to sensitive endpoints',
    value='the system is protected from automated abuse and denial of service attempts',
    criteria=[
        'Anonymous requests: maximum 20 per minute.',
        'Authenticated requests: maximum 100 per minute.',
        'Login endpoint: maximum 60 per minute.',
        'Password reset endpoint: maximum 3 per minute.',
        'Rate-limited responses return HTTP 429 with a Retry-After header.',
    ],
    priority='High', points=3
)

story_block('SEC-004', 'Maintain Audit Log for Inter-Hospital Data Exchanges',
    role='system administrator',
    intent='view an audit log of all WAH4PC gateway transactions',
    value='I can investigate data exchange issues and ensure compliance with data sharing agreements',
    criteria=[
        'Every push and pull transaction to and from the WAH4PC gateway is logged.',
        'Log entry includes transaction type, status, timestamp, patient reference, and raw payload.',
        'Audit log is accessible only to administrators.',
        'Log entries are immutable and cannot be edited or deleted.',
        'Log can be filtered by transaction type, status, and date range.',
    ],
    priority='High', points=5
)

story_block('SEC-005', 'Session Timeout and Token Invalidation',
    role='system administrator',
    intent='enforce a 15-minute idle session timeout',
    value='unattended workstations do not expose patient data',
    criteria=[
        'Access tokens expire after 15 minutes of inactivity.',
        'The frontend prompts the user with a warning 2 minutes before expiry.',
        'Expired sessions redirect to the login page.',
        'Logging out immediately blacklists the current refresh token.',
        'A blacklisted token cannot be used to obtain a new access token.',
    ],
    priority='High', points=5
)

page_break()

# APPENDIX
heading('Appendix A  Story Summary Table', level=1)
divider()
body('The table below lists all 50 user stories with their priority and story point estimates.')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

all_stories = [
    ('AUTH-001','Initiate Staff Account Registration','High','5'),
    ('AUTH-002','Complete Registration via OTP Verification','High','3'),
    ('AUTH-003','Initiate Two-Step Login','High','3'),
    ('AUTH-004','Complete Login via OTP Verification','High','3'),
    ('AUTH-005','Refresh JWT Access Token','High','3'),
    ('AUTH-006','Initiate Password Reset','High','3'),
    ('AUTH-007','Complete Password Reset via OTP','High','3'),
    ('AUTH-008','Change Password While Logged In','Medium','3'),
    ('AUTH-009','Role-Based Access Control','High','8'),
    ('PAT-001','Register a New Patient','High','8'),
    ('PAT-002','Search and View Patient Records','High','5'),
    ('PAT-003','Edit Patient Demographics','Medium','3'),
    ('PAT-004','Record a Patient Condition or Diagnosis','High','5'),
    ('PAT-005','Record Patient Allergies and Intolerances','High','5'),
    ('PAT-006','Record Patient Immunization History','Medium','5'),
    ('PAT-007','Fetch Patient Records from WAH4PC Gateway','High','8'),
    ('PAT-008','Push Patient Records to WAH4PC Gateway','High','8'),
    ('ADM-001','Create a New Encounter','High','5'),
    ('ADM-002','Assign Patient to Ward, Room, and Bed','High','5'),
    ('ADM-003','Record Hospitalization Details','Medium','3'),
    ('ADM-004','Document a Procedure During an Encounter','High','5'),
    ('ADM-005','View and Manage Active Encounters','High','5'),
    ('ADM-006','Close an Encounter','High','3'),
    ('MON-001','Record a Patient Vital Sign Observation','High','5'),
    ('MON-002','Record Multi-Component Observation for Blood Pressure','High','5'),
    ('MON-003','View Patient Observation History','High','5'),
    ('MON-004','Flag Abnormal Observation Values','High','5'),
    ('MON-005','Create a Charge Item for a Billable Monitoring Service','Medium','3'),
    ('LAB-001','Order a Lab Test for a Patient','High','5'),
    ('LAB-002','Record Specimen Collection','High','3'),
    ('LAB-003','Enter and Finalize Lab Test Results','High','8'),
    ('LAB-004','Manage the Lab Test Catalog','Medium','5'),
    ('LAB-005','Track Imaging Studies','Medium','5'),
    ('LAB-006','View Lab Results as a Doctor or Nurse','High','3'),
    ('PHR-001','Prescribe Medication for a Patient','High','8'),
    ('PHR-002','Dispense Medication to a Patient','High','8'),
    ('PHR-003','View and Manage Medication Inventory','High','5'),
    ('PHR-004','Receive Low-Stock Alerts','Medium','3'),
    ('PHR-005','Track Medication Batch Numbers and Expiry Dates','Medium','3'),
    ('PHR-006','View the Prescriptions Queue','High','5'),
    ('BIL-001','Create a Billing Account for an Encounter','High','5'),
    ('BIL-002','Generate an Invoice for a Patient','High','8'),
    ('BIL-003','Submit a PhilHealth Insurance Claim','High','13'),
    ('BIL-004','Record a Payment Against an Invoice','High','5'),
    ('BIL-005','Process Payment Reconciliation','Medium','8'),
    ('BIL-006','Send and View Payment Notices','Low','3'),
    ('DIS-001','Create a Discharge Summary','High','5'),
    ('DIS-002','Record Discharge Instructions and Follow-Up Plan','High','3'),
    ('DIS-003','Finalize and Submit Discharge Record','High','5'),
    ('DIS-004','View Discharge History','Medium','3'),
]

sum_tbl = doc.add_table(rows=len(all_stories)+1, cols=4)
sum_tbl.style = 'Table Grid'

headers = ['Story ID', 'Title', 'Priority', 'Points']
for i, cell in enumerate(sum_tbl.rows[0].cells):
    shade_cell(cell, '0A295C')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.05)
    r = p.add_run(headers[i])
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = WHITE

pri_colors = {'High': 'FFE8E8', 'Medium': 'FFF8E0', 'Low': 'E8F5E9'}
for row_i, (sid, title, pri, pts) in enumerate(all_stories):
    row = sum_tbl.rows[row_i + 1]
    bg  = 'F4F6F8' if row_i % 2 == 0 else 'FFFFFF'
    vals = [sid, title, pri, pts]
    for ci, (cell, val) in enumerate(zip(row.cells, vals)):
        shade_cell(cell, pri_colors.get(pri, bg) if ci == 2 else bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Inches(0.05)
        r = p.add_run(val)
        r.font.name = 'Calibri'
        r.font.size = Pt(8.5)
        r.font.bold = (ci == 0)
        r.font.color.rgb = TEAL if ci == 0 else DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# Footer note
fn = doc.add_paragraph()
fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
rfn = fn.add_run('WAH4H User Stories and System Guidebook  |  Version 1.0  |  June 2026  |  50 Stories Total')
rfn.font.name = 'Calibri'
rfn.font.size = Pt(8.5)
rfn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
rfn.italic = True

doc.save('/workspaces/APC_2025_2026_T1_SS231_G04-WAH-for-Hospitals-WAH4H/docs/WAH4H_User_Stories_Guidebook.docx')
print('User stories document saved.')
