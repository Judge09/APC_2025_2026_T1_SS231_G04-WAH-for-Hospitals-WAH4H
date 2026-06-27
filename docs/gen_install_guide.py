from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

NAVY  = RGBColor(0x0A, 0x29, 0x5C)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
DGRAY = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
ORANGE= RGBColor(0xE6, 0x51, 0x00)
RED   = RGBColor(0xB7, 0x1C, 0x1C)

def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def divider(color='007A87'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading(text, level=1, color=NAVY, size=None):
    sizes = {1: 18, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size if size else sizes[level])
    run.font.color.rgb = color
    return p

def body(text, size=10.5, color=DGRAY, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p

def step_banner(number, title, color_hex='007A87'):
    p = doc.add_paragraph()
    shade_paragraph(p, color_hex)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Inches(0.1)
    run = p.add_run(f'Step {number}:  {title}')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = WHITE
    return p

def code_block(lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, '1E1E2E')
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Inches(0.1)
        run = p.add_run(line)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xA8, 0xFF, 0xC0)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def note_box(text, kind='note'):
    colors  = {'note': 'E3F2FD', 'warning': 'FFF3E0', 'tip': 'E8F5E9'}
    borders = {'note': '1565C0', 'warning': 'E65100', 'tip': '2E7D32'}
    labels  = {'note': 'NOTE', 'warning': 'IMPORTANT', 'tip': 'TIP'}
    label_c = {'note': RGBColor(0x15,0x65,0xC0), 'warning': ORANGE, 'tip': GREEN}

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, colors[kind])

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.left_indent  = Inches(0.1)
    r1 = p1.add_run(labels[kind] + '  ')
    r1.font.name  = 'Calibri'
    r1.font.size  = Pt(9)
    r1.font.bold  = True
    r1.font.color.rgb = label_c[kind]
    r2 = p1.add_run(text)
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(9.5)
    r2.font.color.rgb = DGRAY

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def checklist_item(text, done=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.space_after  = Pt(4)
    mark = 'OK  ' if done else '    '
    run = p.add_run(f'☐  {text}')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10)
    run.font.color.rgb = DGRAY

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
c1 = doc.add_paragraph()
c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(c1, '0A295C')
c1.paragraph_format.space_before = Pt(60)
c1.paragraph_format.space_after  = Pt(4)
r = c1.add_run('WAH4H')
r.font.name = 'Calibri'; r.font.size = Pt(44); r.font.bold = True
r.font.color.rgb = WHITE

c2 = doc.add_paragraph()
c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(c2, '0A295C')
c2.paragraph_format.space_after = Pt(4)
r2 = c2.add_run('We Are Health for Hospitals')
r2.font.name = 'Calibri'; r2.font.size = Pt(16); r2.font.bold = True
r2.font.color.rgb = RGBColor(0xC8, 0x96, 0x20)

c3 = doc.add_paragraph()
c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(c3, '0A295C')
c3.paragraph_format.space_after = Pt(36)
r3 = c3.add_run('Hospital Information System')
r3.font.name = 'Calibri'; r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

sp = doc.add_paragraph()
shade_paragraph(sp, '0A295C')
sp.paragraph_format.space_after = Pt(2)

c4 = doc.add_paragraph()
c4.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(c4, '0A295C')
c4.paragraph_format.space_after = Pt(4)
r4 = c4.add_run('INSTALLATION GUIDE')
r4.font.name = 'Calibri'; r4.font.size = Pt(20); r4.font.bold = True
r4.font.color.rgb = WHITE

c5 = doc.add_paragraph()
c5.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(c5, '0A295C')
c5.paragraph_format.space_after = Pt(4)
r5 = c5.add_run('For System Administrators and Technical Staff')
r5.font.name = 'Calibri'; r5.font.size = Pt(11)
r5.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

c6 = doc.add_paragraph()
c6.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(c6, '0A295C')
c6.paragraph_format.space_after = Pt(60)
r6 = c6.add_run('Version 1.0   |   June 2026')
r6.font.name = 'Calibri'; r6.font.size = Pt(10)
r6.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1  BEFORE YOU START
# ════════════════════════════════════════════════════════════════════════════
heading('1.  Before You Start')
divider()
body(
    'This guide walks you through setting up the WAH4H system on a local machine '
    'or on-premises hospital server. You do not need to be a programmer to follow '
    'these steps, but you do need to be comfortable running commands in a terminal. '
    'Read each step in full before executing it.'
)
note_box(
    'Complete sections 1 and 2 first. Do not skip the prerequisites. '
    'The system will not run if any required software is missing.',
    kind='warning'
)

heading('What You Will Have When Done', level=2)
body('After finishing this guide, you will have the following services running on your machine:')

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, addr, desc) in enumerate([
    ('Header', 'Address', 'What It Is'),
    ('WAH4H Backend API', 'http://localhost:8000', 'The Django server that stores and processes all hospital data.'),
    ('WAH4H Frontend', 'http://localhost:5173', 'The web interface that staff use in their browsers.'),
    ('PostgreSQL Database', 'localhost:5432', 'The database that holds all patient and clinical records.'),
]):
    row = tbl.rows[i]
    if i == 0:
        for cell, txt in zip(row.cells, [label, addr, desc]):
            shade_cell(cell, '0A295C')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Inches(0.05)
            r = p.add_run(txt)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.bold = True
            r.font.color.rgb = WHITE
    else:
        bg = 'F4F6F8' if i % 2 == 0 else 'FFFFFF'
        for ci, (cell, txt) in enumerate(zip(row.cells, [label, addr, desc])):
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Inches(0.05)
            r = p.add_run(txt)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.bold = (ci == 0)
            r.font.color.rgb = TEAL if ci == 0 else DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2  REQUIRED SOFTWARE
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('2.  Required Software')
divider()
body(
    'You must install all four programs below before proceeding. '
    'Each one is free. Download links and installation notes are provided for Windows, '
    'macOS, and Linux.'
)

# 2.1 GIT
heading('2.1  Git', level=2, color=TEAL)
body('Git is used to download the WAH4H source code from the project repository.')

tbl_g = doc.add_table(rows=4, cols=2)
tbl_g.style = 'Table Grid'
rows_g = [
    ('Operating System', 'How to Install'),
    ('Windows', 'Go to https://git-scm.com/download/win and run the installer. Accept all default settings.'),
    ('macOS', 'Open Terminal and run: brew install git'),
    ('Linux (Ubuntu or Debian)', 'Open Terminal and run: sudo apt install git'),
]
for i, (sys_, how) in enumerate(rows_g):
    row = tbl_g.rows[i]
    bg = '0A295C' if i == 0 else ('F4F6F8' if i % 2 == 0 else 'FFFFFF')
    for ci, (cell, txt) in enumerate(zip(row.cells, [sys_, how])):
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Inches(0.05)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5)
        r.font.bold = (i == 0)
        r.font.color.rgb = WHITE if i == 0 else DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('To confirm Git is installed correctly, open a terminal and type:', size=10)
code_block(['git --version', '# You should see something like: git version 2.43.0'])
note_box('On Windows, tick "Add Git to PATH" during installation so you can use it from any terminal.', kind='tip')

# 2.2 Python
heading('2.2  Python 3.12', level=2, color=TEAL)
body(
    'Python is the programming language the WAH4H backend runs on. '
    'You must install version 3.12 exactly. Do not use 3.13 or any version below 3.12.'
)
tbl_py = doc.add_table(rows=4, cols=2)
tbl_py.style = 'Table Grid'
rows_py = [
    ('Operating System', 'How to Install'),
    ('Windows', 'Go to https://www.python.org/downloads and choose Python 3.12.x. Run the installer. On the first screen, tick "Add Python to PATH" before clicking Install Now.'),
    ('macOS', 'Open Terminal and run: brew install python@3.12'),
    ('Linux (Ubuntu or Debian)', 'Open Terminal and run: sudo apt install python3.12 python3.12-venv python3.12-dev'),
]
for i, (sys_, how) in enumerate(rows_py):
    row = tbl_py.rows[i]
    bg  = '0A295C' if i == 0 else ('F4F6F8' if i % 2 == 0 else 'FFFFFF')
    for ci, (cell, txt) in enumerate(zip(row.cells, [sys_, how])):
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent  = Inches(0.05)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.bold = (i == 0)
        r.font.color.rgb = WHITE if i == 0 else DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('To confirm the installation, open a terminal and type:', size=10)
code_block(['python --version', '# Or on Linux/macOS:', 'python3 --version', '# You should see: Python 3.12.x'])

# 2.3 Node.js
heading('2.3  Node.js 20 LTS', level=2, color=TEAL)
body(
    'Node.js runs the WAH4H frontend build process. '
    'You need version 20 LTS (Long-Term Support). '
    'Download it from https://nodejs.org and choose the version labeled "20.x LTS".'
)
body('To confirm the installation:', size=10)
code_block(['node --version', '# You should see: v20.x.x', '', 'npm --version', '# You should see: 10.x.x'])

# 2.4 PostgreSQL
heading('2.4  PostgreSQL 15 or 16', level=2, color=TEAL)
body(
    'PostgreSQL is the database that stores all patient records, clinical data, and system '
    'information. WAH4H requires version 15 or 16.'
)
tbl_pg = doc.add_table(rows=4, cols=2)
tbl_pg.style = 'Table Grid'
rows_pg = [
    ('Operating System', 'How to Install'),
    ('Windows', 'Go to https://www.postgresql.org/download/windows and download the EDB installer. Run it and follow the steps. You will be asked to set a password for the postgres database user. Write this password down. You will need it later.'),
    ('macOS', 'Open Terminal and run: brew install postgresql@15   then start it with:  brew services start postgresql@15'),
    ('Linux (Ubuntu or Debian)', 'Open Terminal and run: sudo apt install postgresql postgresql-contrib   then start it with:  sudo service postgresql start'),
]
for i, (sys_, how) in enumerate(rows_pg):
    row = tbl_pg.rows[i]
    bg  = '0A295C' if i == 0 else ('F4F6F8' if i % 2 == 0 else 'FFFFFF')
    for ci, (cell, txt) in enumerate(zip(row.cells, [sys_, how])):
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent  = Inches(0.05)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.bold = (i == 0)
        r.font.color.rgb = WHITE if i == 0 else DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('To confirm the installation:', size=10)
code_block(['psql --version', '# You should see: psql (PostgreSQL) 15.x'])
note_box(
    'The PostgreSQL installer creates a database administrator account called postgres. '
    'The password you set during installation belongs to this account. Keep it safe. '
    'You will use it in Section 4 of this guide.',
    kind='warning'
)

page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3  DOWNLOAD THE CODE
# ════════════════════════════════════════════════════════════════════════════
heading('3.  Download the Source Code')
divider()
body(
    'Open a terminal window and run the command below. This will download all the '
    'WAH4H project files into a folder on your machine.'
)
code_block([
    'git clone https://github.com/APC-SoCIT/APC_2025_2026_T1_SS231_G04-WAH-for-Hospitals-WAH4H.git',
    '',
    'cd APC_2025_2026_T1_SS231_G04-WAH-for-Hospitals-WAH4H',
])
body('After this command completes, you will see the following folder structure:')
code_block([
    'WAH4H-for-Hospitals/',
    '   wah4h-backend/        <-- Django backend (API server)',
    '   Frontend/',
    '      wah4hospitals-clinic-hub-79-main/   <-- React frontend (web interface)',
    '   deploy/               <-- Docker configuration (optional)',
    '   docs/                 <-- Documentation',
])

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4  BACKEND SETUP
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('4.  Backend Setup (Django API Server)')
divider()
body(
    'The backend is the engine of the system. It stores data, enforces security rules, '
    'and responds to requests from the frontend. Follow each step in order.'
)
note_box(
    'All commands in this section must be run from inside the wah4h-backend folder. '
    'After running the cd command in Step 4.1, stay in that folder for the rest of Section 4.',
    kind='note'
)

# Step 4.1
step_banner('4.1', 'Open the Backend Folder')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('In your terminal, navigate to the backend folder:', indent=0.1)
code_block(['cd wah4h-backend'])

# Step 4.2
step_banner('4.2', 'Create the Database')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('Open the PostgreSQL command line by running:', indent=0.1)
code_block([
    '# On Linux or macOS:',
    'psql -U postgres',
    '',
    '# On Windows (run this in Command Prompt as Administrator):',
    'psql -U postgres',
])
body('You will be asked for the postgres password you set during PostgreSQL installation. After entering it, you will see a prompt that looks like this:   postgres=#', indent=0.1)
body('At that prompt, type the following command exactly and press Enter:', indent=0.1)
code_block(['CREATE DATABASE wah4h_db;'])
body('You should see:   CREATE DATABASE', indent=0.1, italic=True)
body('Then exit PostgreSQL by typing:', indent=0.1)
code_block(['\\q'])

# Step 4.3
step_banner('4.3', 'Create a Python Virtual Environment')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body(
    'A virtual environment keeps the WAH4H backend\'s Python packages separate from other '
    'software on your machine. This prevents conflicts.',
    indent=0.1
)
code_block([
    '# Create the virtual environment:',
    'python -m venv venv',
    '',
    '# Activate it:',
    '# On Windows:',
    'venv\\Scripts\\activate',
    '',
    '# On macOS or Linux:',
    'source venv/bin/activate',
])
body('Your terminal prompt will now show (venv) at the start. This means the virtual environment is active.', indent=0.1, italic=True)
note_box('You must activate the virtual environment every time you start a new terminal session before running the backend.', kind='warning')

# Step 4.4
step_banner('4.4', 'Install Python Packages')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('With the virtual environment active, install all required backend packages:', indent=0.1)
code_block(['pip install -r requirements.txt'])
body('This will download and install Django, the REST framework, database drivers, and all other backend dependencies. It may take one to two minutes.', indent=0.1)

# Step 4.5
step_banner('4.5', 'Create the Configuration File')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body(
    'The backend reads its configuration (database password, secret key, allowed hosts) '
    'from a file named .env. This file is not included in the download for security reasons. '
    'You need to create it from the provided template.',
    indent=0.1
)
code_block([
    '# On macOS or Linux:',
    'cp env.txt .env',
    '',
    '# On Windows:',
    'copy env.txt .env',
])
body('Now open the .env file in any text editor (Notepad works fine on Windows). Fill in the values as shown below:', indent=0.1)

code_block([
    '# Generate your own secret key by running this command first:',
    '# python -c "from django.core.management.utils import get_random_secret_key; print(get_random_secret_key())"',
    '# Copy the output and paste it as the SECRET_KEY value below.',
    '',
    'SECRET_KEY=paste-your-generated-key-here',
    'DEBUG=True',
    'ALLOWED_HOSTS=localhost,127.0.0.1',
    '',
    '# Use the database name you created in Step 4.2',
    'DATABASE_NAME=wah4h_db',
    'DATABASE_USER=postgres',
    'DATABASE_PASSWORD=your-postgres-password-here',
    'DATABASE_HOST=localhost',
    'DATABASE_PORT=5432',
    '',
    '# Gateway settings (leave as-is for local testing)',
    'WAH4PC_GATEWAY_URL=https://wah4pc-gateway.wah.ph',
    'WAH4PC_API_KEY=wah_your-api-key',
    'WAH4PC_PROVIDER_ID=your-provider-uuid',
    'GATEWAY_AUTH_KEY=your-gateway-auth-key',
    '',
    '# Allow the frontend to connect to the backend',
    'CORS_ALLOWED_ORIGINS=http://localhost:5173,http://localhost:3000',
    '',
    'PUBLIC_BASE_URL=http://localhost:8000',
])
note_box(
    'Never share or upload your .env file. It contains your database password and secret key. '
    'It is already excluded from version control by the .gitignore file.',
    kind='warning'
)

page_break()

# Step 4.6
step_banner('4.6', 'Set Up the Database Tables')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body(
    'This command creates all the tables inside the wah4h_db database. '
    'Run it once when setting up, and again whenever the development team releases a system update.',
    indent=0.1
)
code_block(['python manage.py migrate'])
body('You will see a list of operations that end with "OK". This is normal.', indent=0.1, italic=True)

# Step 4.7
step_banner('4.7', 'Load Initial Data')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('Load the hospital organization and laboratory pricing data that the system needs to function:', indent=0.1)
code_block([
    'python manage.py seed_hospitals',
    'python manage.py seed_lab_prices',
])

# Step 4.8
step_banner('4.8', 'Create an Administrator Account')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body(
    'This creates the first login account, which will have full System Administrator access. '
    'You will use these credentials to log in and create accounts for other staff members.',
    indent=0.1
)
code_block(['python manage.py createsuperuser'])
body('The system will ask for a username, email address, and password. Choose a strong password (at least 12 characters).', indent=0.1)

# Step 4.9
step_banner('4.9', 'Start the Backend Server', '1A5276')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('Start the backend API server:', indent=0.1)
code_block(['python manage.py runserver'])
body('You should see output ending with:', indent=0.1)
code_block([
    'Starting development server at http://127.0.0.1:8000/',
    'Quit the server with CTRL-BREAK.',
])
body('The backend is now running. Open a second terminal window for the next section.', indent=0.1, italic=True, color=GREEN)
note_box('Do not close this terminal. The backend will stop if you close or interrupt it.', kind='warning')

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5  FRONTEND SETUP
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('5.  Frontend Setup (Web Interface)')
divider()
body(
    'The frontend is the visual part of the system that staff see in their web browsers. '
    'Open a new terminal window for these steps. Leave the backend terminal from Section 4 running.'
)
note_box('All commands in this section must be run from inside the Frontend/wah4hospitals-clinic-hub-79-main folder.', kind='note')

step_banner('5.1', 'Open the Frontend Folder')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('Navigate to the frontend folder from the project root:', indent=0.1)
code_block(['cd Frontend/wah4hospitals-clinic-hub-79-main'])

step_banner('5.2', 'Install Frontend Packages')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('Download and install all the JavaScript packages the frontend needs:', indent=0.1)
code_block(['npm install'])
body('This may take one to three minutes depending on your internet connection.', indent=0.1, italic=True)

step_banner('5.3', 'Create the Frontend Configuration File')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body(
    'The frontend needs to know the address of the backend server. '
    'Create a file named .env inside the frontend folder and paste the following content into it:',
    indent=0.1
)
code_block([
    'LOCAL_8000=http://localhost:8000',
    'STURDY_ADVENTURE_BASE=http://localhost:8000',
    'STURDY_ADVENTURE_BASE_8000=http://localhost:8000',
    'BACKEND_PHARMACY=http://localhost:8000/api/pharmacy/',
    'BACKEND_PHARMACY_8000=http://localhost:8000/api/pharmacy/',
    'BACKEND_ADMISSIONS=http://localhost:8000/api/admissions/',
    'BACKEND_ADMISSIONS_8000=http://localhost:8000/api/admissions/',
    'BACKEND_BILLING=http://localhost:8000/api/billing/',
    'BACKEND_BILLING_8000=http://localhost:8000/api/billing/',
    'BACKEND_DISCHARGE=http://localhost:8000/api/discharge/',
    'BACKEND_DISCHARGE_8000=http://localhost:8000/api/discharge/',
    'BACKEND_ACCOUNTS=http://localhost:8000/accounts/',
    'BACKEND_ACCOUNTS_8000=http://localhost:8000/accounts/',
    'BACKEND_PATIENTS=http://localhost:8000/api/patients/',
    'BACKEND_PATIENTS_8000=http://localhost:8000/api/patients/',
    'BACKEND_MONITORING=http://localhost:8000/api/',
    'BACKEND_MONITORING_8000=http://localhost:8000/api/',
    'BACKEND_LABORATORY=http://localhost:8000/api/laboratory/',
    'BACKEND_LABORATORY_8000=http://localhost:8000/api/laboratory/',
])

step_banner('5.4', 'Start the Frontend Server', '1A5276')
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body('Start the frontend development server:', indent=0.1)
code_block(['npm run dev'])
body('You should see output that includes:', indent=0.1)
code_block([
    'VITE v5.x.x  ready in xxx ms',
    '',
    '  Local:   http://localhost:5173/',
])
body('The frontend is now running. Open your web browser and go to http://localhost:5173', indent=0.1, italic=True, color=GREEN)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6  FIRST LOGIN
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('6.  First Login and System Verification')
divider()
body(
    'With both the backend (Section 4) and frontend (Section 5) running, '
    'open a web browser and go to this address:'
)
body('http://localhost:5173', size=14, color=TEAL)
body('Use the administrator account you created in Step 4.8 to log in.')
note_box('The login page uses a two-step verification process. After entering your email and password, check your email for a one-time password (OTP) and enter it on the next screen.', kind='note')

heading('Verification Checklist', level=2)
body('After logging in, confirm each of the following to make sure the installation is working correctly:')

checks = [
    'The login page loads at http://localhost:5173',
    'You can log in with the administrator account created in Step 4.8',
    'The dashboard is visible after a successful login',
    'The backend health check at http://localhost:8000/api/health/ returns a status of ok',
    'The patient list page loads without errors',
    'The navigation menu shows all modules (Patients, Admission, Monitoring, Laboratory, Pharmacy, Billing, Discharge)',
]
for c in checks:
    checklist_item(c)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
note_box(
    'If the health check at http://localhost:8000/api/health/ shows any red or failed services, '
    'check that PostgreSQL is running and that the database credentials in .env are correct.',
    kind='warning'
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7  STOPPING AND RESTARTING
# ════════════════════════════════════════════════════════════════════════════
heading('7.  Stopping and Restarting the System')
divider()
body('To stop the system, press Ctrl + C in each terminal window (backend and frontend) and then close them.')
body('To restart the system on a subsequent day:')

code_block([
    '# Terminal 1: Restart the backend',
    'cd wah4h-backend',
    'source venv/bin/activate        # Windows: venv\\Scripts\\activate',
    'python manage.py runserver',
    '',
    '# Terminal 2: Restart the frontend',
    'cd Frontend/wah4hospitals-clinic-hub-79-main',
    'npm run dev',
])
note_box('You do not need to repeat Sections 3, 4.2 through 4.8 again. Those steps only need to be done once during the initial setup.', kind='tip')

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8  TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('8.  Troubleshooting')
divider()
body('This section covers the most common problems encountered during installation and how to fix them.')

def trouble_block(problem, cause, fix_lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, 'FFF8F8')

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.left_indent  = Inches(0.1)
    r_label = p1.add_run('PROBLEM  ')
    r_label.font.name = 'Calibri'; r_label.font.size = Pt(9); r_label.font.bold = True
    r_label.font.color.rgb = RED
    r_title = p1.add_run(problem)
    r_title.font.name = 'Calibri'; r_title.font.size = Pt(10.5); r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p2 = cell.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.1)
    p2.paragraph_format.space_after = Pt(2)
    rc = p2.add_run('Likely cause: ')
    rc.font.name = 'Calibri'; rc.font.size = Pt(9.5); rc.font.bold = True
    rc.font.color.rgb = DGRAY
    rc2 = p2.add_run(cause)
    rc2.font.name = 'Calibri'; rc2.font.size = Pt(9.5)
    rc2.font.color.rgb = DGRAY

    pf = cell.add_paragraph()
    pf.paragraph_format.left_indent = Inches(0.1)
    pf.paragraph_format.space_after = Pt(3)
    rf = pf.add_run('How to fix:')
    rf.font.name = 'Calibri'; rf.font.size = Pt(9.5); rf.font.bold = True
    rf.font.color.rgb = GREEN

    for line in fix_lines:
        pl = cell.add_paragraph()
        pl.paragraph_format.left_indent = Inches(0.2)
        pl.paragraph_format.space_before = Pt(1)
        pl.paragraph_format.space_after  = Pt(2)
        rl = pl.add_run(line)
        rl.font.name = 'Calibri'; rl.font.size = Pt(9.5)
        rl.font.color.rgb = DGRAY

    cell.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

trouble_block(
    'connection refused when running python manage.py migrate',
    'PostgreSQL is not running or the credentials in .env are wrong.',
    [
        '1. Confirm PostgreSQL is running.',
        '   Linux: sudo service postgresql status',
        '   macOS: brew services list | grep postgresql',
        '   Windows: Open Services from the Start menu and look for postgresql-x64-15.',
        '2. Open the .env file and check that DATABASE_PASSWORD matches the password you set when installing PostgreSQL.',
        '3. Confirm that DATABASE_HOST is set to localhost and DATABASE_PORT is set to 5432.',
    ]
)

trouble_block(
    'pip install fails with error about psycopg2',
    'The C++ compiler required to build the PostgreSQL adapter is missing (Windows only).',
    [
        '1. Download Microsoft C++ Build Tools from https://visualstudio.microsoft.com/visual-cpp-build-tools/',
        '2. Run the installer and select "Desktop development with C++" from the workload list.',
        '3. Click Install and wait for it to finish.',
        '4. Close and reopen your terminal, then run pip install -r requirements.txt again.',
    ]
)

trouble_block(
    'The login page shows a CORS error in the browser developer console',
    'The frontend is trying to connect to a backend address that is not in the allowed origins list.',
    [
        '1. Open the .env file inside wah4h-backend.',
        '2. Find the CORS_ALLOWED_ORIGINS line.',
        '3. Make sure it includes http://localhost:5173',
        '   Example: CORS_ALLOWED_ORIGINS=http://localhost:5173,http://localhost:3000',
        '4. Save the file and restart the backend server (Ctrl+C, then python manage.py runserver).',
    ]
)

trouble_block(
    'npm install fails with a message about an unsupported Node.js version',
    'The Node.js version on your machine is not version 20.',
    [
        '1. Check your version: node --version',
        '2. If it is not v20.x.x, install Node.js 20 LTS from https://nodejs.org',
        '3. If you use nvm (Node Version Manager), run: nvm install 20  and then  nvm use 20',
        '4. Run npm install again.',
    ]
)

trouble_block(
    'Port 8000 or 5173 is already in use',
    'Another program or a previous server process is occupying that port.',
    [
        'On Linux or macOS, find and stop the process:',
        '   lsof -ti:8000 | xargs kill -9',
        '   lsof -ti:5173 | xargs kill -9',
        '',
        'On Windows (run Command Prompt as Administrator):',
        '   netstat -ano | findstr :8000',
        '   Look at the PID number in the last column, then run:',
        '   taskkill /PID <that-number> /F',
        '',
        'Then try starting the server again.',
    ]
)

trouble_block(
    'python manage.py migrate fails with "table already exists"',
    'The database is in an inconsistent state, usually from a failed previous setup.',
    [
        '1. Drop and recreate the database (this will delete all data):',
        '   psql -U postgres -c "DROP DATABASE wah4h_db;"',
        '   psql -U postgres -c "CREATE DATABASE wah4h_db;"',
        '2. Run migrations again:',
        '   python manage.py migrate',
        '3. Re-seed the data:',
        '   python manage.py seed_hospitals',
        '   python manage.py seed_lab_prices',
        '4. Recreate the administrator account:',
        '   python manage.py createsuperuser',
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9  QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════════
page_break()
heading('9.  Quick Reference Card')
divider()
body('Print or save this page as a daily reference for starting and stopping the system.')

tbl_qr = doc.add_table(rows=2, cols=1)
tbl_qr.style = 'Table Grid'

c_start = tbl_qr.rows[0].cells[0]
shade_cell(c_start, 'E8F5E9')
ps = c_start.paragraphs[0]
ps.paragraph_format.space_before = Pt(6)
ps.paragraph_format.left_indent  = Inches(0.1)
rs = ps.add_run('Starting the System (run once per day)')
rs.font.name = 'Calibri'; rs.font.size = Pt(11); rs.font.bold = True; rs.font.color.rgb = GREEN

start_cmds = [
    '# Terminal 1 (Backend):',
    'cd wah4h-backend',
    'source venv/bin/activate     # Windows: venv\\Scripts\\activate',
    'python manage.py runserver',
    '',
    '# Terminal 2 (Frontend):',
    'cd Frontend/wah4hospitals-clinic-hub-79-main',
    'npm run dev',
    '',
    '# Then open a browser and go to:',
    'http://localhost:5173',
]
for cmd in start_cmds:
    pc = c_start.add_paragraph()
    pc.paragraph_format.left_indent = Inches(0.1)
    pc.paragraph_format.space_before = Pt(1)
    pc.paragraph_format.space_after  = Pt(1)
    rc = pc.add_run(cmd)
    rc.font.name = 'Courier New'; rc.font.size = Pt(9)
    rc.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

c_start.add_paragraph().paragraph_format.space_after = Pt(4)

c_stop = tbl_qr.rows[1].cells[0]
shade_cell(c_stop, 'FFF3E0')
pt = c_stop.paragraphs[0]
pt.paragraph_format.space_before = Pt(6)
pt.paragraph_format.left_indent  = Inches(0.1)
rt = pt.add_run('Stopping the System')
rt.font.name = 'Calibri'; rt.font.size = Pt(11); rt.font.bold = True; rt.font.color.rgb = ORANGE
stop_p = c_stop.add_paragraph()
stop_p.paragraph_format.left_indent = Inches(0.1)
stop_p.paragraph_format.space_after = Pt(6)
rsp = stop_p.add_run('Press Ctrl + C in both terminal windows.')
rsp.font.name = 'Calibri'; rsp.font.size = Pt(10); rsp.font.color.rgb = DGRAY

doc.add_paragraph().paragraph_format.space_after = Pt(8)

heading('Key Addresses', level=2)
addr_tbl = doc.add_table(rows=4, cols=2)
addr_tbl.style = 'Table Grid'
for i, (label, val) in enumerate([
    ('Service', 'Address'),
    ('WAH4H Web Interface (Frontend)', 'http://localhost:5173'),
    ('WAH4H API Server (Backend)', 'http://localhost:8000'),
    ('System Health Check', 'http://localhost:8000/api/health/'),
]):
    row = addr_tbl.rows[i]
    bg  = '0A295C' if i == 0 else ('F4F6F8' if i % 2 == 0 else 'FFFFFF')
    for ci, (cell, txt) in enumerate(zip(row.cells, [label, val])):
        shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent  = Inches(0.05)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.font.bold = (i == 0 or ci == 0)
        r.font.color.rgb = WHITE if i == 0 else (TEAL if ci == 1 else DGRAY)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

fn = doc.add_paragraph()
fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
rfn = fn.add_run('WAH4H Installation Guide  |  Version 1.0  |  June 2026')
rfn.font.name = 'Calibri'; rfn.font.size = Pt(8.5)
rfn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
rfn.italic = True

doc.save('/workspaces/APC_2025_2026_T1_SS231_G04-WAH-for-Hospitals-WAH4H/docs/WAH4H_Installation_Guide.docx')
print('Installation guide saved.')
